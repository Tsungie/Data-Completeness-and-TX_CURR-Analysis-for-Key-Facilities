from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_cv():
    document = Document()

    # Set margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Styles
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # --- HEADER ---
    header_color = RGBColor(0, 51, 102) # Navy Blue
    accent_color = RGBColor(0, 128, 128) # Teal

    name_paragraph = document.add_paragraph()
    name_run = name_paragraph.add_run('TSUNGIE NCUBE')
    name_run.bold = True
    name_run.font.size = Pt(24)
    name_run.font.color.rgb = header_color
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_paragraph = document.add_paragraph()
    title_run = title_paragraph.add_run('BUSINESS ANALYST | DATA SCIENTIST | MBA CANDIDATE (AI)')
    title_run.bold = True
    title_run.font.size = Pt(12)
    title_run.font.color.rgb = accent_color
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contact_paragraph = document.add_paragraph()
    contact_paragraph.add_run('Harare, Zimbabwe | +263 776 844 533 | misstsungie@gmail.com').bold = False
    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    link_paragraph = document.add_paragraph()
    link_run = link_paragraph.add_run('linkedin.com/in/tsungie-ncube-a627bb62')
    link_run.font.color.rgb = RGBColor(0, 0, 255)
    link_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph() # Spacer

    # --- FUNCTION TO ADD SECTION HEADERS ---
    def add_section_header(text):
        p = document.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Shading for the background (Teal bar)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:val'), 'clear')
        shading_elm.set(qn('w:color'), 'auto')
        shading_elm.set(qn('w:fill'), '008080') # Teal hex
        p._p.get_or_add_pPr().append(shading_elm)

    # --- PROFESSIONAL SUMMARY ---
    add_section_header('Professional Summary')
    summary = document.add_paragraph()
    summary.add_run("Strategic Data-Driven Professional with 9 years of experience bridging the gap between technical software development and business strategy. Currently pursuing an MBA in Artificial Intelligence, combining deep technical expertise in Python, SQL, and Data Visualization with strong stakeholder management skills. Proven track record of analyzing complex datasets to drive operational efficiency and implementing large-scale systems across 900+ facilities. Expert in translating complex data into actionable business stories for non-technical leadership.")
    summary.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- TECHNICAL & BUSINESS TOOLKIT ---
    add_section_header('Technical & Business Toolkit')
    
    # Create a table for skills
    table = document.add_table(rows=1, cols=2)
    table.autofit = True
    
    # Cell 1
    cell1 = table.cell(0, 0)
    c1_p = cell1.paragraphs[0]
    c1_p.add_run("Data Analysis & Visualization:").bold = True
    c1_p.add_run("\nPython (Pandas, NumPy), SQL (Advanced), Tableau, Streamlit, OpenRefine, Excel (VBA).")
    
    # Cell 2
    cell2 = table.cell(0, 1)
    c2_p = cell2.paragraphs[0]
    c2_p.add_run("Business Analysis & Strategy:").bold = True
    c2_p.add_run("\nRequirements Gathering, Stakeholder Management, Agile/Scrum, Predictive Modeling, KPI Tracking.")

    document.add_paragraph() # Spacer

    # --- PROFESSIONAL EXPERIENCE ---
    add_section_header('Professional Experience')

    # JOB 1
    p_job1 = document.add_paragraph()
    p_job1.add_run('Lead Product Manager & IT Team Lead').bold = True
    p_job1.add_run(' | ZIMTTECH').bold = False
    p_job1.add_run('\nNov 2022 – Nov 2025 | Harare, Zimbabwe').italic = True
    
    ul_job1 = document.add_paragraph()
    ul_job1.style = 'List Bullet'
    ul_job1.add_run("Business Analysis & Strategy: Led the end-to-end product strategy for a nationwide EHR platform serving 900+ facilities. Gathered requirements from Ministry of Health officials and clinical stakeholders to define roadmaps and validate solutions.")
    
    ul_job1_b = document.add_paragraph()
    ul_job1_b.style = 'List Bullet'
    ul_job1_b.add_run("Data Analysis & Predictive Modeling: Analyzed longitudinal treatment data for 10,000+ ART patients using Python. Contributed insights for a predictive model to identify patient default risks, enabling proactive intervention.")
    
    ul_job1_c = document.add_paragraph()
    ul_job1_c.style = 'List Bullet'
    ul_job1_c.add_run("Process Automation: Integrated systems using FHIR standards, reducing lab result turnaround times from hours to minutes and eliminating manual data transfer errors.")
    
    ul_job1_d = document.add_paragraph()
    ul_job1_d.style = 'List Bullet'
    ul_job1_d.add_run("Reporting & KPIs: Designed and maintained dashboards to track system adoption and clinical outcomes, presenting regular progress reports to executive leadership.")

    # JOB 2
    p_job2 = document.add_paragraph()
    p_job2.add_run('Provincial IT Manager & Product Lead').bold = True
    p_job2.add_run(' | Columbia University (ICAP)').bold = False
    p_job2.add_run('\nFeb 2020 – Oct 2022 | Harare, Zimbabwe').italic = True
    
    ul_job2 = document.add_paragraph()
    ul_job2.style = 'List Bullet'
    ul_job2.add_run("System Integration: Orchestrated complex integrations between clinical systems and national health databases, ensuring full data consistency and HIPAA compliance.")
    
    ul_job2_b = document.add_paragraph()
    ul_job2_b.style = 'List Bullet'
    ul_job2_b.add_run("Stakeholder Engagement: Managed IT strategy for HIV programs across 8 districts. Acted as the primary liaison between technical teams and clinical staff to ensure technology solutions met business needs.")
    
    ul_job2_c = document.add_paragraph()
    ul_job2_c.style = 'List Bullet'
    ul_job2_c.add_run("Training & Adoption: Developed training materials and conducted workshops for 500+ staff, ensuring successful rollout and high user adoption rates.")

    # JOB 3
    p_job3 = document.add_paragraph()
    p_job3.add_run('Business Analyst & Software Engineer').bold = True
    p_job3.add_run(' | Research Triangle International').bold = False
    p_job3.add_run('\nFeb 2016 – Nov 2018 | Harare, Zimbabwe').italic = True
    
    ul_job3 = document.add_paragraph()
    ul_job3.style = 'List Bullet'
    ul_job3.add_run("Database Design & Reporting: Designed database architectures managing 500,000+ records. Built custom analytics and reporting solutions (JasperReports/SQL) to support ministry-level decision-making.")
    
    ul_job3_b = document.add_paragraph()
    ul_job3_b.style = 'List Bullet'
    ul_job3_b.add_run("Requirements Engineering: Conducted user research with researchers and clinical staff to translate complex workflows into intuitive software modules.")

    document.add_paragraph() # Spacer

    # --- EDUCATION ---
    add_section_header('Education & Certifications')

    edu_p1 = document.add_paragraph()
    edu_p1.add_run("Master of Business Administration (MBA) in Artificial Intelligence").bold = True
    edu_p1.add_run("\nCumbria University, UK | 2024 – 2026 (In Progress)")
    
    edu_p2 = document.add_paragraph()
    edu_p2.add_run("Certificate in Data Science and Machine Learning").bold = True
    edu_p2.add_run("\nMIT Institute for Data, Systems, and Society (IDSS) | 2022 – 2023")
    
    edu_p3 = document.add_paragraph()
    edu_p3.add_run("BSc (Honors) Software Engineering with Multimedia").bold = True
    edu_p3.add_run("\nLimkokwing University, Botswana | 2008 – 2012")

    return document

# Generate and save
doc = create_cv()
file_path = "Tsungie_Ncube_Etihad_CV.docx"
doc.save(file_path)

file_path